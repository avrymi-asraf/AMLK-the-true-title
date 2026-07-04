"""Build docs/failure-examples.docx (and .md) from error-analysis JSON.

Run from repo root: python scripts/build_failure_examples.py
Requires: python-docx (uv pip install python-docx)
"""
import json
import re
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
OUT_MD = ROOT / "docs/failure-examples.md"
OUT_DOCX = ROOT / "docs/failure-examples.docx"

# Lines that introduce a following text block (article / prediction / reference).
BLOCK_LABELS = {
    "Article (full):",
    "Model prediction (full):",
    "Reference summary (full):",
}

# Top-level section headings (not examples).
SECTION_HEADINGS = {
    "Summary failure rates (n=100 each)",
    "Multi-header format (pipe-separated digests)",
    "Reference summaries (ground truth)",
    "Finetuned v3 predictions",
    "Finetuned v3 — example failures",
    "Base v3 (zero-shot, raw prompt — known baseline bug)",
}


def has_hebrew(text: str, n: int = 100) -> bool:
    return any("\u05d0" <= c <= "\u05ea" for c in text[:n])


def pick(examples: list[dict], label_set: set[str]) -> dict | None:
    for e in examples:
        if set(e["labels"]) == label_set:
            return e
    return None


def plain(s: str) -> str:
    """Strip markdown bold/backticks for paste-friendly prose."""
    s = re.sub(r"\*\*([^*]+)\*\*", r"\1", s)
    s = re.sub(r"`([^`]+)`", r"\1", s)
    return s


CLUSTER_DESC = {
    "entity_or_number_error+hallucination+omission": (
        "Hallucination + wrong entities + missing main points (20% of finetuned sample). "
        "The model writes a plausible Hebrew media-digest in the right format, but invents newspapers, "
        "people, and claims; it also drops key stories from the article."
    ),
    "hallucination+omission": (
        "Hallucination + omission (15%). The summary is on-topic in tone but describes events "
        "that are not in the article and skips the article's actual main points."
    ),
    "entity_or_number_error+omission": (
        "Wrong entities + omission (10%). Names/outlets are garbled or wrong; main content missing."
    ),
    "entity_or_number_error+hallucination+lead_copying": (
        "Hallucination + entity errors + lead bias (8%). Picks up lead-adjacent names but fabricates the rest."
    ),
    "lead_copying+omission": (
        "Lead copying + omission (7%). Echoes opening themes or entities but fails to summarize the full article."
    ),
    "entity_or_number_error+fluency_problem+hallucination": (
        "Hallucination + entity errors + fluency (7%). Fluent-looking Hebrew with garbled tokens and invented content."
    ),
    "omission": (
        "Omission only (8%). May look reasonable but leaves out the article's central story."
    ),
    "wrong_language": (
        "Wrong language / non-answer (97% of base sample, post-hoc tag). "
        "Qwen3 enters an English thinking block and never produces a Hebrew summary. "
        "The literature failure taxonomy marks these as clean because the English reasoning restates the article."
    ),
    "entity_or_number_error+hallucination": (
        "Hallucination + entity errors in English (13% of base). "
        "Responds in English with invented facts (e.g. dollar amounts not in the source)."
    ),
    "entity_or_number_error+hallucination+omission_base": (
        "English summary with wrong facts and missing points (11% of base)."
    ),
}


def write_text_block(lines: list[str], label: str, text: str) -> None:
    lines.append(label)
    lines.append("")
    lines.append(text)
    lines.append("")


def write_example(lines: list[str], e: dict, combo: tuple[str, ...], key: str) -> None:
    label_str = ", ".join(combo) if combo else "(none — wrong_language tagged post-hoc)"
    lines.append(f"Failure labels: {label_str}")
    lines.append("")
    lines.append(plain(CLUSTER_DESC[key]))
    lines.append("")
    write_text_block(lines, "Article (full):", e["text"].strip())
    write_text_block(lines, "Model prediction (full):", e["prediction"].strip())
    write_text_block(lines, "Reference summary (full):", e["reference"].strip())
    lines.append("What went wrong:")
    bullets: list[str] = []
    if not combo or key == "wrong_language":
        bullets += [
            "Model outputs English reasoning inside a thinking block and often hits the token limit before any Hebrew summary.",
            "Judge taxonomy has no wrong-language label, so this often scores as no failure.",
        ]
    else:
        if "hallucination" in combo:
            bullets.append("Prediction invents outlets, people, or events not supported by the article.")
        if "omission" in combo:
            bullets.append("Key stories in the reference (and article) are missing from the prediction.")
        if "entity_or_number_error" in combo:
            bullets.append("Names/numbers are wrong or garbled (e.g. mixed Hebrew/Latin characters).")
        if "lead_copying" in combo:
            bullets.append("Over-relies on lead-adjacent entities without faithful abstraction.")
        if "fluency_problem" in combo:
            bullets.append("Surface fluency breaks down (tokenization artifacts, incoherent phrases).")
        if not has_hebrew(e["prediction"]):
            bullets.append("Responds in English, not Hebrew.")
    for b in bullets:
        lines.append(f"• {b}")
    lines.append("")


def multi_header_stats(texts: list[str]) -> dict:
    def seg_count(text: str) -> int:
        return len([s for s in text.split("|") if s.strip()])

    n = len(texts)
    multi = sum(1 for t in texts if seg_count(t) > 1)
    has_pipe = sum(1 for t in texts if "|" in t)
    dist: dict[int, int] = {}
    for t in texts:
        c = seg_count(t)
        dist[c] = dist.get(c, 0) + 1
    return {"n": n, "multi": multi, "has_pipe": has_pipe, "dist": dist}


def append_multi_header_section(lines: list[str], preds: list[str]) -> None:
    pred = multi_header_stats(preds)
    lines += [
        "Multi-header format (pipe-separated digests)",
        "",
        "HeSum references often use a multi-header style: several short headline clauses joined by "
        '" | " (a weekly media-roundup digest, not a single-topic abstract). A summary counts as '
        "multi-header when it has more than one non-empty segment after splitting on |.",
        "",
        "Reference summaries (ground truth)",
        "",
        "• Test set (n=1000): 254 multi-header (25.4%), 746 single-segment (74.6%)",
        "• Train set (n=8000): 2068 multi-header (25.9%), 5932 single-segment (74.2%)",
        "",
        "Test-set segment distribution:",
        "• 1 segment: 746 (74.6%)",
        "• 2 segments: 16 (1.6%)",
        "• 3 segments: 205 (20.5%)",
        "• 4 segments: 27 (2.7%)",
        "• 5 segments: 5 (0.5%)",
        "• 6 segments: 1 (0.1%)",
        "",
        "Most multi-header references have exactly 3 pipe-separated items (205/254 = 81% of multi-header refs).",
        "",
        "Finetuned v3 predictions",
        "",
        f"• Contains | separator: {pred['has_pipe']} ({100 * pred['has_pipe'] / pred['n']:.1f}%)",
        f"• Multi-header (>1 segment): {pred['multi']} ({100 * pred['multi'] / pred['n']:.1f}%)",
        f"• Single segment: {pred['n'] - pred['multi']} ({100 * (pred['n'] - pred['multi']) / pred['n']:.1f}%)",
        "",
        "The model over-uses the pipe template: about 62% of predictions are multi-header vs 25% of references "
        "(more than 2× the reference rate). It learned the HeSum digest format but applies it too often and runs too long:",
        "",
        "• Reference median length: ~147 chars / ~25 words",
        "• Prediction median length: ~548 chars / ~89 words (3.7× reference)",
        f"• Predictions reach up to {max(pred['dist'])} pipe segments in one output; references almost never exceed 6",
        "",
        "Interpretation: pipe-headline overfitting is a distinct failure mode from hallucination. The model produces "
        'plausible-looking "| newspaper X does Y" clauses even when the reference is a single-topic summary (74.6% of refs). '
        "This inflates format correctness while hurting faithfulness and length control.",
        "",
    ]


def main() -> None:
    ft = json.loads((ROOT / "outputs/results/finetuned-v3.errors.json").read_text())
    base = json.loads((ROOT / "outputs/results/base-v3.errors.json").read_text())
    ft_ex, base_ex = ft["examples"], base["examples"]

    lines = [
        "Failure Examples — v3 (Qwen3-2B, whole variant)",
        "",
        "AMLK Hebrew news summarization project. Curated examples of the biggest failure modes in the v3 "
        "evaluation run (3-epoch LoRA, anti-degeneration decode). Failure labels from Gemini gemini-2.5-flash-lite "
        "on a 100-example sample. Full machine-readable reports: outputs/results/finetuned-v3.errors.json and "
        "outputs/results/base-v3.errors.json.",
        "",
        "Summary failure rates (n=100 each)",
        "",
    ]
    for k in ["hallucination", "omission", "entity_or_number_error", "lead_copying", "fluency_problem"]:
        name = k.replace("_", " ")
        lines.append(f"• {name}: finetuned {ft['failure_rates'][k]:.0%}, base {base['failure_rates'][k]:.0%}")
    lines += [
        "• wrong language (post-hoc tag): finetuned 0%, base 97%",
        "",
        "Finetuned: fails on content (hallucination/omission) but always answers in Hebrew.",
        "Base: fails on task completion (English thinking, no Hebrew answer) unless the chat template is used.",
        "",
    ]
    preds = [
        json.loads(line)["prediction"]
        for line in (ROOT / "outputs/results/predictions-finetuned.jsonl").read_text().splitlines()
        if line.strip()
    ]
    append_multi_header_section(lines, preds)

    lines += [
        "Finetuned v3 — example failures",
        "",
    ]

    ft_clusters = [
        (("entity_or_number_error", "hallucination", "omission"), "entity_or_number_error+hallucination+omission", "1"),
        (("hallucination", "omission"), "hallucination+omission", "2"),
        (("entity_or_number_error", "omission"), "entity_or_number_error+omission", "3"),
        (("entity_or_number_error", "hallucination", "lead_copying"), "entity_or_number_error+hallucination+lead_copying", "4"),
        (("lead_copying", "omission"), "lead_copying+omission", "5"),
        (("entity_or_number_error", "fluency_problem", "hallucination"), "entity_or_number_error+fluency_problem+hallucination", "6"),
        (("omission",), "omission", "7"),
    ]
    for combo, key, num in ft_clusters:
        e = pick(ft_ex, set(combo))
        if not e:
            continue
        lines.append(f"Example {num}: {key.replace('_', ' ')}")
        lines.append("")
        write_example(lines, e, combo, key)
        lines.append("")

    lines += [
        "Base v3 (zero-shot, raw prompt — known baseline bug)",
        "",
        "Note: a fairer base baseline uses the Qwen3 chat template with thinking disabled. "
        "Regen job 6a48d621 (--pred-suffix=-v4) was submitted 2026-07-04.",
        "",
    ]

    e = next(x for x in base_ex if not x["labels"] and not has_hebrew(x["prediction"]))
    lines.append("Example 1: wrong language (non-answer)")
    lines.append("")
    write_example(lines, e, tuple(), "wrong_language")
    lines.append("")

    for combo, key, num in [
        (("entity_or_number_error", "hallucination"), "entity_or_number_error+hallucination", "2"),
        (("entity_or_number_error", "hallucination", "omission"), "entity_or_number_error+hallucination+omission_base", "3"),
    ]:
        e = pick(base_ex, set(combo))
        if not e:
            continue
        lines.append(f"Example {num}: {key.replace('_', ' ')}")
        lines.append("")
        write_example(lines, e, combo, key)
        lines.append("")

    OUT_MD.write_text("\n".join(lines), encoding="utf-8")
    write_docx(lines, OUT_DOCX)
    print(f"Wrote {OUT_MD} ({len(lines)} lines)")
    print(f"Wrote {OUT_DOCX}")


def _set_rtl(paragraph) -> None:
    """Right-align and mark paragraph RTL for Hebrew text."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement

    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_pr = paragraph._element.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    p_pr.append(bidi)


def write_docx(lines: list[str], path: Path) -> None:
    """Render the plain-text report as a Word file (opens in Google Docs via upload)."""
    from docx import Document

    doc = Document()
    title = lines[0] if lines else "Failure Examples"
    doc.add_heading(title, level=0)

    i = 1
    while i < len(lines):
        line = lines[i]
        if not line.strip():
            i += 1
            continue

        if line.startswith("Example "):
            doc.add_heading(line, level=2)
            i += 1
            continue

        if line in SECTION_HEADINGS:
            doc.add_heading(line, level=1)
            i += 1
            continue

        if line in BLOCK_LABELS:
            p = doc.add_paragraph()
            p.add_run(line).bold = True
            i += 1
            # Collect block until blank line or next structural line.
            block: list[str] = []
            while i < len(lines):
                nxt = lines[i]
                if not nxt.strip():
                    break
                if nxt in BLOCK_LABELS or nxt.startswith("Example ") or nxt in SECTION_HEADINGS:
                    break
                if nxt == "What went wrong:":
                    break
                block.append(nxt)
                i += 1
            text = "\n".join(block).strip()
            if text:
                body = doc.add_paragraph(text)
                if has_hebrew(text):
                    _set_rtl(body)
            continue

        if line == "What went wrong:":
            p = doc.add_paragraph()
            p.add_run(line).bold = True
            i += 1
            while i < len(lines) and lines[i].strip().startswith("•"):
                doc.add_paragraph(lines[i].strip()[2:].strip(), style="List Bullet")
                i += 1
            continue

        if line.startswith("•"):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
            i += 1
            continue

        if line.startswith("Failure labels:"):
            p = doc.add_paragraph()
            p.add_run(line).bold = True
            i += 1
            continue

        # Body paragraph (intro, interpretation, notes).
        body = doc.add_paragraph(line)
        if has_hebrew(line):
            _set_rtl(body)
        i += 1

    doc.save(path)


if __name__ == "__main__":
    main()
